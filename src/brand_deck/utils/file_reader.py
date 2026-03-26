# file_reader.py — 多格式文件读取器
"""Read various input formats and normalize to plain text."""

from __future__ import annotations

import os
from pathlib import Path


def read_file(path: str | Path) -> str:
    """
    Read a file and return its text content.
    Supports: .txt, .md, .docx, .pdf
    For unsupported formats, raises ValueError with supported format list.
    """
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")

    suffix = path.suffix.lower()

    if suffix in (".txt", ".md", ".markdown", ".csv"):
        return _read_text(path)
    elif suffix == ".docx":
        return _read_docx(path)
    elif suffix == ".pdf":
        return _read_pdf(path)
    else:
        raise ValueError(
            f"Unsupported file format: {suffix}\n"
            f"Supported formats: .txt, .md, .docx, .pdf"
        )


def read_multiple(paths: list[str | Path]) -> str:
    """Read multiple files and concatenate their text content."""
    parts = []
    for p in paths:
        p = Path(p)
        if p.is_dir():
            # Read all supported files in directory
            for child in sorted(p.iterdir()):
                if child.suffix.lower() in (".txt", ".md", ".docx", ".pdf"):
                    try:
                        parts.append(f"--- {child.name} ---\n{read_file(child)}")
                    except Exception as e:
                        parts.append(f"--- {child.name} --- [Error: {e}]")
        else:
            parts.append(f"--- {p.name} ---\n{read_file(p)}")
    return "\n\n".join(parts)


def _read_text(path: Path) -> str:
    """Read plain text file with encoding fallback."""
    for encoding in ("utf-8", "utf-8-sig", "latin-1", "gbk"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError(f"Cannot decode file: {path} (tried utf-8, latin-1, gbk)")


def _read_docx(path: Path) -> str:
    """Read .docx file using python-docx."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx is required to read .docx files.\n"
            "Install it: pip install python-docx"
        )

    doc = Document(str(path))
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    return "\n".join(paragraphs)


def _read_pdf(path: Path) -> str:
    """Read .pdf file using pdfplumber."""
    try:
        import pdfplumber
    except ImportError:
        raise ImportError(
            "pdfplumber is required to read .pdf files.\n"
            "Install it: pip install pdfplumber"
        )

    text_parts = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)

    return "\n\n".join(text_parts)
